import os.path
import argparse
import re
import shutil
import requests
from requests.auth import HTTPBasicAuth
from docx.oxml import OxmlElement, ns
import sys
from docx import Document
from docx.shared import Pt
import subprocess
from docx.shared import Cm
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from GitLabPipelineInfo import GitLabPipelineInfo
from datetime import datetime,timedelta
import subprocess
import sys
from docx.oxml.ns import qn  
import json
from configparser import ConfigParser
import xml.etree.ElementTree as ET
import datetime
import glob


class Logger(object):

    def __init__(self, filename='default.log', stream=sys.stdout):
        self.terminal = stream
        self.log = open(filename, 'a',encoding='utf-8')
    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)
    def flush(self):
        pass


def shell(command, result=False):
    process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE,shell=True)
    stdout, stderr = process.communicate()
    return_code = process.wait()
    if result:
        if return_code != 0:
            print(command,': 命令执行失败', return_code)

    return return_code, stdout.decode()


def get_base(url):
    sub_path = url.replace('.git', '')
    sub_path = os.path.basename(sub_path)
    print(sub_path)
    return sub_path


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)

    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def get_cell_text(cell):
    """安全获取单元格文本（兼容超链接）"""
    text = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                text.append(run.text)
    return "".join(text)


def get_common_commit(branch1, branch2):
    print(branch1,branch2)
    """
    使用git merge-base命令直接查找两个分支的共同祖先提交
    """
    try:
        cmd = f"git merge-base {branch1} {branch2}"
        result = subprocess.run(cmd, shell=True, check=True,
                              stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                              text=True)
        return result.stdout.strip()
    except subprocess.CalledProcessError as e:
        print(f"错误: {e}")
        exit(1)

def deal_description(source_tag,target_tag):
    print(source_tag,target_tag)
    change_shell = 'git log {}..{} --format="%H|%ci|%an|%s|%b"'.format(source_tag,target_tag)
    _, change_android = shell(change_shell)

    if change_android == "":
        print(os.getcwd(),"==========================没有修改，取最新提交==========================")
        cmd = 'git log {} --format="%H|%ci|%an|%s|%b" -1'.format(target_tag)
        _,change_android = shell(cmd)
    
    log_all = change_android.split('\n')
    log_del = []
    for i in log_all:
        if '|' in i:
            sp = i.split('|')
            str_s = sp[3]
            str_b = sp[4]

            if 'Merge branch' in str_s:
                des = i.replace(f'{str_s}|{str_b}', str_b)
            else:
                des = i.replace(f'{str_s}|{str_b}', str_s)
                if len(str_b)==0:
                    des = des[:-1]
            log_del.append(des)
    return log_del


# 得到最新提交
def get_latest_remote_commit(repo_url, vcpu_tag=''):
    if repo_url.endswith("civic-vcpu-rsu3.0.git") or repo_url.endswith("civic-vcpu-rsu3.5.git"):
        print("----------------------VCPU no tag------------------------------")
        # vcpu target tag
        vcpu_target_tag = vcpu_tag.split('..')[1]
        return vcpu_target_tag
    else:
        cmd = f"git ls-remote {repo_url} HEAD"
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        
        if result.returncode != 0:
            raise RuntimeError(f"git ls-remote failed: {result.stderr}")
        if not result.stdout.strip():
            raise ValueError(f"No output from git ls-remote for {repo_url}")
        
        return result.stdout.split()[0]


def get_change_data(source_tag, target_tag,
                   url_and='https://git.gitlab.com/sb/android/aosp-manifest.git', vcpu_tag='', sed=True,
                   down=True):
    # 保存当前工作目录
    original_path = os.getcwd()
    main_path_and = get_base(url_and)
    sub_path = os.path.join(os.getcwd(), main_path_and)
    print("sub_path:", sub_path)

    try:
        if down:
            if os.path.exists(sub_path):
                shutil.rmtree(sub_path)
        
        # 下载代码
        if not os.path.exists(sub_path):
            print("init cloning............", url_and)
            shell('git clone {}'.format(url_and))
            # 进入目录进行其他git操作
            os.chdir(sub_path)
            shell('git fetch -p')
            shell('git checkout {}'.format(source_tag))
        else:
            os.chdir(sub_path)
            
        print("*"*100)
        print("当前工作目录:", os.getcwd())
        print("*"*100)
        
        # 处理VCPU特殊情况
        if len(vcpu_tag) > 0:
            print('---------------------------VCPU----------------------')
            split1 = vcpu_tag.split('..')
            if len(split1) == 2:
                source_tag = split1[0]
                target_tag = split1[1]
            else:
                print('分支出错', target_tag, url_and)
                exit(1)
                
        # 检查sorce_tag标签是否存在
        check_cmd = f"git ls-remote --tags {url_and} {source_tag}"
        result = subprocess.run(check_cmd, shell=True, capture_output=True, text=True)
        if not result.stdout.strip():
            print(f"警告: source_tag '{source_tag}' 不存在！")
            print(f"将使用传入的commitID: {url_and} {source_tag}")

        # 检查target_tag标签是否存在
        check_cmd = f"git ls-remote --tags {url_and} {target_tag}"
        result = subprocess.run(check_cmd, shell=True, capture_output=True, text=True)
        if not result.stdout.strip():
            print(f"警告: target_tag '{target_tag}' 不存在！")
            target_tag = get_latest_remote_commit(url_and, vcpu_tag)
            print(f"将使用最新提交: {url_and} {target_tag}")

        # 获取共同祖先提交
        hax_des1 = get_common_commit(source_tag, target_tag)

        # 切换到目标标签
        shell(f'git checkout {target_tag}')

        # 处理变更描述
        change_android = deal_description(source_tag, target_tag)
        if change_android:  # 检查是否为空
            hash_new = change_android[0].split('|')[0]
        else:
            hash_new = target_tag  # 如果没有变更，使用目标标签作为hash
            
        change_android2 = []
        if sed:
            for i, j in enumerate(change_android):
                sub_data = change_android[i]
                change_android2 = change_android[:i]
        else:
            change_android2 = change_android
            
        return hash_new, change_android2
        
    except Exception as e:
        print(f"在get_change_data函数中发生错误: {e}")
        raise
    finally:
        # 确保无论发生什么都会切换回原始目录
        os.chdir(original_path)
        print(f"已切换回原始目录: {original_path}")

def get_table0(table0):
    android_version = ''
    qnx_version = ''
    vcpu_version = ''
    source_branch = ''
    for i in table0.rows:
        cell = i.cells[0]
        version_data = cell.text
        if 'E_RELEASE' in version_data:
            source_branch = i.cells[1].text
        if 'ANDROID_REVISION' in version_data:
            android_version = i.cells[1].text
        if 'QNX_REVISION' in version_data:
            qnx_version = i.cells[1].text
        if 'VCPU_REVISION' in version_data:
            vcpu_version = i.cells[1].text
    return source_branch, android_version, qnx_version, vcpu_version


def set_copy_data(table0):
    for j, i in enumerate(table0.rows):      
        first_row = table0.rows[0]
        source_cell = first_row.cells[1]
        target_cell = first_row.cells[2]

        target_cell.text = ''
            
        if source_cell.paragraphs:
            hyperlinks = source_cell.paragraphs[0]._element.xpath('.//w:hyperlink')
            if hyperlinks:
                hyperlink = hyperlinks[0]
                r_id = hyperlink.get(qn('r:id'))
                hyperlink_url = source_cell.part.rels[r_id]._target
                build_number = hyperlink_url.split('/')[-1]

                source_paragraph = source_cell.paragraphs[0]
                target_paragraph = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph()
                    
                add_hyperlink(target_paragraph, build_number, hyperlink_url)
            else:
                target_cell.text = source_cell.text
                print("复制纯文本内容")
        i.cells[2].text = i.cells[1].text
    set_width1(table0)


def set_columns1(table, list_params, url):
    for i, row in enumerate(table.rows):
        if i == 0:
            row.cells[1].paragraphs[0].clear()
            add_hyperlink(row.cells[1].paragraphs[0], list_params[i], url)
        else:
            row.cells[1].text = list_params[i]

def delete_rows_except_first(table):
    for i, row in enumerate(table.rows):
        if i != 0:
            row._element.getparent().remove(row._element)


def deal_auther(name, upper):
    if upper:
        name = name.strip().upper()
    else:
        name = name.strip()
    return name


def star_jiry(d):
    if d.startswith('APRICOT') or "APRICOT" in d:
        list_jiry = re.findall('(APRICOT.*?\\d+)', d)

        if len(list_jiry)==0:
            x1 = 'No Jira-ID Found'
        else:
            x1 = list_jiry[0]

    else:
        x1 = 'No Jira-ID Found'
    return x1


def deal_url(url):
    url = url.replace('git.gitlab.com','git.gitlab.com')
    return url

def set_word_data(table, data, url, upper=True):
    delete_rows_except_first(table)

    for index, string in enumerate(data):
        if len(string.strip()) == 0:
            continue
        hash, daily, auther, message = string.split('|')
        b = hash[:8]
        a = deal_auther(auther, upper=upper)
        d = message.strip()
        x1 = ''
        url = deal_url(url)

        row_cells = table.add_row().cells

        url1 = url.replace('.git', '/-/commit/')
        url1 = url1 + b
        add_hyperlink(row_cells[0].paragraphs[0], b, url1)
        row_cells[1].text = a
        row_cells[2].text = d

        x1 = star_jiry(d)

        row_cells[3].text = x1

    set_width(table)


def set_word_data2(table, data, upper=True):
    delete_rows_except_first(table)

    for index, string in enumerate(data):
        hash, auther,time_x, url, message = string
        bx1 = hash
        ax1 = deal_auther(auther,upper=upper)
        dx1 = str(message.strip())
        url = deal_url(url)       

        row_cells = table.add_row().cells

        add_hyperlink(row_cells[0].paragraphs[0], bx1, url)
        row_cells[1].text = ax1
        try:
            row_cells[2].text = dx1
        except:
            print('\n\n出错了----error')

        x1 = star_jiry(dx1)

        row_cells[3].text = x1

    set_width(table)

def time_deal(ts):
    if isinstance(ts, str) and not ts.isdigit():
        try:
            dt_utc = datetime.datetime.strptime(ts, "%Y-%m-%d %H:%M:%S.%f")
            beijing_time = dt_utc + timedelta(hours=8)
        except ValueError:
            raise ValueError("Invalid time string format. Expected format: 'YYYY-MM-DD HH:MM:SS.ffffff'")
    else:
        timestamp = int(ts) / 1000
        dt_utc = datetime.utcfromtimestamp(timestamp)
        beijing_time = dt_utc + timedelta(hours=8)

    formatted_time = beijing_time.strftime("%b %d, %Y, %I:%M:%S %p")
    return formatted_time


def cr_dir(path_name, remove=True):
    if not os.path.exists(path_name):
        print('没有文件夹，创建', path_name)
        os.mkdir(path_name)
    else:
        if remove:
            shutil.rmtree(path_name)
            os.mkdir(path_name)


def get_xml(artifactory_url, token, source_xml,target_xml):
    # if not os.path.exists("manifest"):
    #     os.mkdir("manifest")

    down_command = "jf rt dl \"{}\" \"{}\" --url={} --user=GAYUXIA --password={} --flat=true".format(source_xml,target_xml,artifactory_url,token)
    status,result = shell(down_command,result=True)


def parse_manifest_xml(xml_file):
    """
    解析 manifest XML 文件，返回项目字典：
    {项目路径: (项目名称, 版本号, 远程URL)}

    逻辑参考 _parse_manifest：
    - 根据 remote.fetch 判断是否拼接路径前缀（如 sb/android/AOSP）
    - 保留完整返回结构 (name, revision, full_url)
    """
    projects = {}
    default_base = "https://git.gitlab.com/"

    try:
        tree = ET.parse(xml_file)
        root = tree.getroot()

        # 获取 remote.fetch
        fetch = ""
        for remote in root.findall("remote"):
            fetch_raw = remote.get("fetch", "")
            fetch = fetch_raw.replace(default_base, "").strip()
            # 只取第一个 remote 即可（一般 manifest 只有一个 fetch）
            break

        for project in root.findall("project"):
            name = project.get("name", "").strip()
            path = project.get("path", name).strip()
            revision = project.get("revision", "main").strip()

            # 根据 fetch 规则拼接完整仓库路径
            if fetch == "sb/android/AOSP":
                repo_name = os.path.join(fetch, name).replace("\\", "/")
            else:
                repo_name = name

            # 构建完整远程 URL
            full_url = f"{default_base}{repo_name}.git"

            # ⚙️ 保留原始返回结构
            projects[path] = (repo_name, revision, full_url)

    except Exception as e:
        print(f"解析 manifest 文件失败: {e}")

    return projects


def get_gitlab_commits(project_path, private_token, gitlab_url, old_revision, new_revision):
    """通过GitLab API获取两个版本之间的提交记录"""
    commits = []
    try:
        # 编码项目路径
        encoded_project_path = requests.utils.quote(project_path, safe='')
        
        # GitLab API URL
        api_url = f"{gitlab_url}/api/v4/projects/{encoded_project_path}/repository/compare"
        
        headers = {
            'Private-Token': private_token
        }
        
        params = {
            'from': old_revision,
            'to': new_revision
        }
        # eg: https://gitlab.example.com/api/v4/projects/my-group%2Fmy-project/repository/compare?from=old_revision&to=new_revision
        response = requests.get(api_url, headers=headers, params=params, timeout=30)
        
        if response.status_code == 200:
            data = response.json()
            for commit in data.get('commits', []):
                commit_info = {
                    'hash': commit['id'][:8],
                    'author': commit['author_name'],
                    'time': commit['created_at'],
                    'message': commit['title'],
                    'url': commit['web_url']
                }
                commits.append(commit_info)
        else:
            print(f"GitLab API请求失败: {response.status_code} - {response.text}")
            
    except Exception as e:
        print(f"获取GitLab提交记录失败: {e}")
        
    return commits


def compare_manifests_and_get_commits(old_manifest, new_manifest, private_token, gitlab_url):
    """对比两个manifest文件并获取所有变更的提交记录"""
    all_commits = []
    
    # 解析manifest文件
    old_projects = parse_manifest_xml(old_manifest)
    new_projects = parse_manifest_xml(new_manifest)
    
    print(f"旧manifest项目数: {len(old_projects)}")
    print(f"新manifest项目数: {len(new_projects)}")
    
    # 对比项目版本变化
    for path, (name, new_revision, url) in new_projects.items():
        if path in old_projects:
            old_name, old_revision, old_url = old_projects[path]
            
            # 如果版本号不同，获取提交记录
            if old_revision != new_revision:
                print(f"检测到变更: {name} ({path})")
                print(f"  旧版本: {old_revision} -> 新版本: {new_revision}")
                
                commits = get_gitlab_commits(name, private_token, gitlab_url, old_revision, new_revision)
                all_commits.extend(commits)
                
                print(f"  找到 {len(commits)} 个提交")
        else:
            # 新添加的项目
            print(f"新项目: {name} ({path}) - 版本: {new_revision}")
            
            # 获取该项目的所有提交（从初始提交到当前版本）
            commits = get_gitlab_commits(name, private_token, gitlab_url, 'initial', new_revision)
            all_commits.extend(commits)
            
            print(f"  找到 {len(commits)} 个提交")
    
    # 检查被删除的项目
    for path in old_projects:
        if path not in new_projects:
            name, revision, url = old_projects[path]
            print(f"被删除的项目: {name} ({path}) - 版本: {revision}")
    
    return all_commits


def format_commit_data(commits):
    """格式化提交数据为表格需要的格式"""
    formatted_data = []
    for commit in commits:
        formatted_data.append([
            commit['hash'],
            commit['author'],
            commit['time'],
            commit['url'],
            commit['message']
        ])
    return formatted_data


def set_align(table):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_width1(table):
    table.autofit = False
    for i in range(3):
        if i == 0:
            xx = 3.46
        else:
            xx = 7.71
        for j in range(len(table.rows)):
            table.rows[j].cells[i].width = Cm(xx)
            if i == 0:
                table.rows[j].height = Cm(1.02)

def set_font_paragraph(paragraph,font_size=10):
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        
def set_font(table,font_size=10):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing_rule = 0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                set_font_paragraph(paragraph,font_size)

def set_width(table):
    a = 2.06
    b = 3.18
    c = 10
    d = 3.18

    all = a + b + c + d
    for i in range(4):
        if i == 0:
            xx = a
        elif i == 1:
            xx = b
        elif i == 2:
            xx = c
        elif i == 3:
            xx = d
        for j in range(len(table.rows)):
            table.rows[j].cells[i].width = Cm(xx)
            if i == 0:
                table.rows[j].height = Cm(0.51)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing_rule = 0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)


if __name__ == '__main__':
    main_path = os.getcwd()
    os.chdir(main_path)
    # 删除之前的log .docx .xml
    patterns = ["*.log", "*.docx", "*.xml"]
    for pattern in patterns:
        for file_path in glob.glob(pattern):
            os.remove(file_path)

    # 加载环境变量
    config = ConfigParser()
    config.read("/home/gayuxia/config.ini")
    PRIVATE_TOKEN = config.get("gitlab","token")
    GITLAB_URL = config.get("gitlab","url")
    GROUP_PATH = "sb"

    # 下载快照xml
    artifactory_url = config.get("artifact","url")
    artifactory_token = config.get("artifact","token")
    user = config.get("artifact","username")


    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")  # 获取当前时间戳
    sys.stdout = Logger(f'{timestamp}.log')
    print('开启日志记录')
    # root_dir = os.getcwd()

    # 创建 ArgumentParser 对象
    parser = argparse.ArgumentParser(description='Process some integers.')

    # 定义命名参数
    parser.add_argument('--build_release', default='E069.0', help='Release version')
    parser.add_argument('--gitlab_build_url',
                        default='https://git.gitlab.com/sb/trigger/sb-system-staging/-/pipelines/464642',
                        help='gitlab url')
    parser.add_argument('--source_vendor_xml', help="source venrodr manifest xml")
    parser.add_argument('--target_vendor_xml', default='vendor_manifest.xml', help="target vendor manfiest xml")
    parser.add_argument('--source_qssi_xml', help="source qssi manifest xml")
    parser.add_argument('--target_qssi_xml', default='qssi_manifest.xml', help="target qssi manifest xml")
    parser.add_argument('--source_qnx_xml', help="source qnx manfiest xml")
    parser.add_argument('--target_qnx_xml', default='manifest.xml', help="target qnx manfiest xml")
    parser.add_argument('--android_url', default='https://git.gitlab.com/sb/android/aosp-manifest.git',
                        help='Android manifest git project')
    parser.add_argument('--qnx_url', default='https://git.gitlab.com/sb/qnx/qnx-manifest.git',
                        help='QNX manifest git project')
    parser.add_argument('--amss_url', default='https://git.gitlab.com/sb/sa8295p-hqx-4-5-5-1_amss_standard_oem.git',
                        help='amss manifest git project')
    parser.add_argument('--vcpu_url', default='https://git.sb.com/apricotqal/civic-vcpu-rsu3.0.git',
                        help='Vcpu manifest git project')
    parser.add_argument('--vcpu_tag', default='E067.0_Delivery_RSU..E068.0_Delivery_RSU',
                        help="Vcpu'shell that to get git log")

    parser.add_argument('--source_tag', default='E228.0-19.4', help='source version tag')
    parser.add_argument('--target_tag', default='E228.0-19.4', help='target version tag') 
    parser.add_argument('--changelog_type', default='STAR3.0', help='STAR3.0 OR STAR3.5')                 
    parser.add_argument('--word_url', default='https://artifact.gitlab.com/ui/native/sb/sb/sb_Temporary/GAYUXIA/changelog/E241.0_STAR3.0_changelog.docx', help="use jfrog url or local file: --word_url E068.0_Changelog.docx")

    # 解析命令行参数
    args = parser.parse_args()
    build_release = args.build_release
    source_vendor_xml = args.source_vendor_xml
    target_vendor_xml = args.target_vendor_xml
    source_qssi_xml = args.source_qssi_xml
    target_qssi_xml = args.target_qssi_xml
    source_qnx_xml = args.source_qnx_xml
    target_qnx_xml = args.target_qnx_xml
    changelog_type = args.changelog_type

    source_tag = args.source_tag
    target_tag = args.target_tag

    build_url = args.gitlab_build_url

    des_branch = build_release

    android_url = args.android_url
    qnx_url = args.qnx_url
    vcpu_url = args.vcpu_url
    amss_url = args.amss_url

    vendor_xml = '8295_master_vendor.xml'
    qssi_xml = '8295_master_qssi.xml'
    qnx_xml = '8295_master.xml'

    vcpu_tag = args.vcpu_tag

    word_url = args.word_url
    if 'http' in word_url:
        path_docx = os.path.basename(word_url)
    else:
        path_docx = word_url

    # 下载changelog模板文件
    if not os.path.exists(os.path.join(os.getcwd(),path_docx)):
        down_command=f'curl -L -H "X-JFrog-Art-Api: {artifactory_token}" "{word_url}" -o "{path_docx}" --fail'
        shell(down_command)
        os.system(down_command)

        
    # 得到 gitlab build number, 和构建时间
    pipeline_tool = GitLabPipelineInfo(
            private_token = PRIVATE_TOKEN
            )
    build_number = pipeline_tool.extract_pipeline_id(build_url)

    if build_number:
        created_time, updated_time = pipeline_tool.get_pipeline_time(
                project_path = "sb/trigger/sb-system-staging",
                pipeline_id = build_number
        )
        
    ts = str(created_time)
    print(f"传入的时间戳 ts = {ts}") 
    build_timestamp = time_deal(ts)

    des_docx = path_docx.replace(path_docx.split('_')[0], build_release)
    original_path_docx = os.getcwd()
    path_docx1 = os.path.join(original_path_docx, path_docx)
    print('original word abs path==',path_docx1,'==')
    print('生成的文档：',des_docx)

    doc = Document(path_docx1)

    # 获取表一数据
    temp_list = get_table0(doc.tables[0])

    source_branch=temp_list[0]
    android_version=temp_list[1]
    qnx_version=temp_list[2]
    vcpu_version=temp_list[3]

    # 把第一个表格中第二列的内容，复制到第三列中
    set_copy_data(doc.tables[0])  
 
    # 得到aosp 侧manifest仓库的changelog
    hash_android, android_data = get_change_data(source_tag, target_tag, android_url)
 
    # 得到qnx侧manifest仓库的changelog
    hash_qnx, qnx_data = get_change_data(source_tag, target_tag, qnx_url)
    
    # 得到amss侧仓库的changelog
    hash_amss, amss_data = get_change_data(source_tag, target_tag, amss_url)
    
    # 得到VCPU 仓库上一个版本到当前版本的changelog
    hash_vcpu, vcpu_data = get_change_data(source_tag, target_tag, vcpu_url, vcpu_tag=vcpu_tag, sed=False)
    print('VCPU-------------------------------------------------\n' * 3)

    set_columns1(doc.tables[0], [build_number, build_timestamp, build_release, hash_android, hash_qnx, hash_vcpu],build_url)

    set_word_data(doc.tables[1], android_data, android_url)
    set_word_data(doc.tables[2], qnx_data, qnx_url)
    set_word_data(doc.tables[5], amss_data, amss_url)
    set_word_data(doc.tables[6], vcpu_data, vcpu_url, True)
    

    # 下载source manifest快照文件
    get_xml(artifactory_url, artifactory_token, source_vendor_xml, "source_vendor_manifest.xml")
    get_xml(artifactory_url, artifactory_token, source_qssi_xml, "surce_qssi_manifest.xml")
    get_xml(artifactory_url, artifactory_token, source_qnx_xml, "source_qnx_manifest.xml")
    # 下载target manifest快照文件
    get_xml(artifactory_url, artifactory_token, target_vendor_xml, "target_vendor_manifest.xml")
    get_xml(artifactory_url, artifactory_token, target_qssi_xml, "target_qssi_manifest.xml")
    get_xml(artifactory_url, artifactory_token, target_qnx_xml, "target_qnx_manifest.xml")

    print('开始通过manifest对比分析变更...\n' * 3)
    
    # Android侧变更分析 (Vendor + QSSI)
    print("分析Android侧变更...")
    android_commits = []
    
    if os.path.exists("target_vendor_manifest.xml"):
        vendor_commits = compare_manifests_and_get_commits("source_vendor_manifest.xml", "target_vendor_manifest.xml", PRIVATE_TOKEN, GITLAB_URL)
        android_commits.extend(vendor_commits)
        print(f"Vendor侧找到 {len(vendor_commits)} 个提交")
    else:
        print("\033[1;32m vendor manifest文件不存在，跳过vendor分析 \033[0m")
    
    if os.path.exists("target_qssi_manifest.xml"):
        qssi_commits = compare_manifests_and_get_commits("surce_qssi_manifest.xml", "target_qssi_manifest.xml", PRIVATE_TOKEN, GITLAB_URL)
        android_commits.extend(qssi_commits)
        print(f"QSSI侧找到 {len(qssi_commits)} 个提交")
    else:
        print("\033[1;32m qssi manifest文件不存在，跳过qssi分析 \033[0m")

    # 格式化Android提交数据
    data_android_raw = format_commit_data(android_commits)

    # === ✅ Android changelog 去重逻辑开始 ===
    # 使用 (hash, message) 作为唯一键，避免重复提交
    seen = set()
    data_android = []
    for commit in data_android_raw:
        # commit结构: [hash, author, time, url, message]
        key = (commit[0], commit[4])
        if key not in seen:
            seen.add(key)
            data_android.append(commit)

    print(f"去重后 Android 提交数量: {len(data_android_raw)} → {len(data_android)}")
    # === ✅ Android changelog 去重逻辑结束 ===

    # 写入 android changelog.txt
    changelog_txt = os.path.join(main_path, 'android')
    cr_dir(changelog_txt)
    log_android_txt = os.path.join(changelog_txt, 'changelog.txt')
    with open(log_android_txt, 'w', encoding='utf-8') as f:
        for commit in data_android:
            f.write('|'.join(commit) + '\n')

    # 更新 Word 文档中的 Android 表格
    set_word_data2(doc.tables[3], data_android)

    # QNX侧变更分析
    print("分析QNX侧变更...")
    if os.path.exists("target_qnx_manifest.xml"):
        qnx_commits = compare_manifests_and_get_commits("source_qnx_manifest.xml", "target_qnx_manifest.xml", PRIVATE_TOKEN, GITLAB_URL)
        data_qnx = format_commit_data(qnx_commits)
        print(f"QNX侧找到 {len(qnx_commits)} 个提交")
    else:
        data_qnx = []
        print("\033[1;32m QNX manifest文件不存在，跳过QNX分析 \033[0m")

    # 写入QNX changelog.txt
    changlog_txt = os.path.join(main_path,'qnx')
    cr_dir(changlog_txt)
    log_qnx_txt = os.path.join(changlog_txt, 'changelog.txt') 
    with open(log_qnx_txt, 'w', encoding='utf-8') as f:
        for commit in data_qnx:
            f.write('|'.join(commit) + '\n')

    # 更新Word文档中的QNX表格
    set_word_data2(doc.tables[4], data_qnx)



    # 设置表格格式
    for i in range(6):
        set_align(doc.tables[i])
    set_width1(doc.tables[0])
    set_font(doc.tables[0])
    
    # 保存文档
    doc.save('change.docx')
    shutil.copy('change.docx',des_docx)
    print('生成的文档：',des_docx)

    # 上传生成的文档
    word_generate = args.word_url.replace(path_docx.split('_')[0], build_release)
    print("开始上传changlog到jfrog!!!")
    upload_command=f'curl -O -H "X-JFrog-Art-Api: {artifactory_token}" -T {des_docx} {word_generate}'
    shell(upload_command, result=True)
   
    print('\n完成')